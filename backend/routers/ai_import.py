"""AI test importer — convert an uploaded test (PDF / Word / image / text) into
the site's exam format, optionally alongside a separate answer sheet that the
model must treat as the authoritative source of correct answers.

Providers (selected with the LLM_PROVIDER env var, default "gemini"):
- "gemini" — Google Gemini via the google-genai SDK. Free online option and
  multimodal: PDFs and images are sent as raw bytes (Gemini reads charts,
  scanned pages and layout natively, so we do NOT pre-extract text). Needs
  GEMINI_API_KEY; GEMINI_MODEL defaults to "gemini-2.5-flash".
- "claude" — Anthropic Claude via tool-use. Images and PDFs are sent as base64
  (PDFs as native document blocks, so scans work); other files have their text
  extracted locally first. Needs ANTHROPIC_API_KEY.
- "local" — no external API. Extracts text locally and returns it as a single
  reading section for the teacher to finish by hand. Useful offline / for tests.

Design notes:
- Heavy/optional deps (google-genai, anthropic, pypdf, python-docx) are imported
  lazily inside the handler so the app still boots before
  `pip install -r backend/requirements.txt`.
- Without the relevant API key the endpoint returns 503 so the frontend can show
  a friendly "set up AI import" message.
- Every provider returns the same JSON shape (the build_ielts_test schema) so the
  frontend visual builder is unchanged.
"""
import os
import io
import time
import json
import base64
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from backend.service.database import get_db
from backend.service import models
from backend.service.auth_deps import require_role
from backend.service.subskills import SUB_SKILLS
from backend.service.ratelimit import rate_limit

router = APIRouter(prefix="/api/import", tags=["import"])

# Guard the paid LLM import: 10 uploads/min per user.
_ai_limiter = rate_limit(10, 60)

# Anthropic model for the `claude` import provider. Override with ANTHROPIC_MODEL.
MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-5")

# Cap the uploaded test file so a huge upload can't exhaust server memory. We
# read at most MAX_UPLOAD_BYTES + 1 bytes, so an oversized file is rejected
# without ever being fully loaded.
MAX_UPLOAD_BYTES = 20 * 1024 * 1024  # 20 MB

# Full IELTS tests are long. Too small an output-token ceiling truncates the JSON
# and the parse fails, so give the models plenty of room (overridable via env).
GEMINI_MAX_OUTPUT_TOKENS = int(os.getenv("GEMINI_MAX_OUTPUT_TOKENS", "65536"))
CLAUDE_MAX_OUTPUT_TOKENS = int(os.getenv("CLAUDE_MAX_OUTPUT_TOKENS", "16000"))
# Per-request timeout so a hung LLM call can't tie up a worker indefinitely.
LLM_TIMEOUT_SECONDS = float(os.getenv("LLM_TIMEOUT_SECONDS", "120"))
# Substrings that mark a retryable/transient upstream error.
_TRANSIENT_KEYS = (
    "503", "529", "unavailable", "overloaded", "high demand",
    "429", "rate limit", "resource_exhausted", "timeout", "deadline",
)

# Tool schema the model must fill — mirrors tests_io.TestIn so the result can be
# saved directly through POST /api/tests/import after teacher review.
TEST_TOOL = {
    "name": "build_ielts_test",
    "description": "Return the IELTS test parsed from the supplied material.",
    "input_schema": {
        "type": "object",
        "properties": {
            "name": {"type": "string", "description": "Title of the whole test"},
            "difficulty": {"type": "string", "enum": ["low", "medium", "high"]},
            "time_limit_min": {"type": "integer"},
            "source_question_count": {"type": "integer", "description": "The highest question number printed in the source material (official IELTS numbering — a 'Choose TWO letters' task covering questions 24-25 counts as TWO). A full IELTS paper is 40."},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "skill": {"type": "string", "enum": ["reading", "listening", "writing", "speaking"]},
                        "title": {"type": "string"},
                        "passage_md": {"type": "string", "description": "All student-facing material for the section in markdown: the reading passage or listening transcript PLUS task instructions, any list of headings, word/option banks, example answers and notes/table/flow-chart templates. Never omit shared lists students must choose from."},
                        "questions": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "qtype": {"type": "string", "enum": ["mcq", "short", "explain"]},
                                    "qformat": {"type": "string", "enum": ["tfng", "ynng", "matching", "multi_select", "gap_fill"], "description": "IELTS display format: tfng=True/False/Not Given, ynng=Yes/No/Not Given, matching=pick from a shared list, multi_select=choose N letters, gap_fill=complete the blank. Omit for plain multiple choice / short answer."},
                                    "prompt": {"type": "string"},
                                    "task_instructions": {"type": "string", "description": "The exact instruction block printed above this question's task, including word limits (e.g. 'Complete the notes below. Write ONE WORD ONLY from the passage for each answer.'). Use the IDENTICAL string for every question of the same task."},
                                    "options": {"type": "array", "items": {"type": "string"}, "description": "MCQ options"},
                                    "correct_index": {"type": "integer", "description": "0-based index of the correct MCQ option"},
                                    "correct_indices": {"type": "array", "items": {"type": "integer"}, "description": "multi_select only: 0-based indices of ALL correct options"},
                                    "select_count": {"type": "integer", "description": "multi_select only: how many options the student must choose"},
                                    "accept_answers": {"type": "array", "items": {"type": "string"}, "description": "Accepted answers for short questions"},
                                    "sub_skill": {"type": "string", "enum": SUB_SKILLS, "description": "Question category from the fixed list"},
                                    "explanation": {"type": "string", "description": "For reading/listening questions: 2-3 plain-language sentences on why the correct answer is correct, briefly noting why a common wrong choice is a trap."},
                                    "support_sentences": {"type": "array", "items": {"type": "string"}, "description": "For reading/listening questions: the exact sentence(s) copied verbatim from passage_md that justify the answer."},
                                    "paraphrases": {
                                        "type": "array",
                                        "description": "For reading/listening questions: 1-3 wording pairs showing how the question rephrases the passage.",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "question_phrase": {"type": "string", "description": "Short phrase from the question or correct option"},
                                                "passage_phrase": {"type": "string", "description": "The corresponding phrase copied verbatim from passage_md"},
                                            },
                                            "required": ["question_phrase", "passage_phrase"],
                                        },
                                    },
                                },
                                "required": ["qtype", "prompt"],
                            },
                        },
                    },
                    "required": ["skill", "title", "questions"],
                },
            },
        },
        "required": ["name", "sections"],
    },
}

SYSTEM = (
    "You convert raw IELTS test material into a structured test. Be precise and "
    "consistent — extract exactly what is on the page, never invent content.\n"
    "Rules:\n"
    "1. One section per reading passage or listening recording. Put the full passage "
    "or listening transcript verbatim in passage_md. For writing/speaking sections, "
    "put the task instructions (and any chart description) in passage_md.\n"
    "2. Set each section's skill to one of: reading, listening, writing, speaking.\n"
    "3. Question types (qtype) and display formats (qformat) — match the official "
    "computer-based IELTS exactly:\n"
    "   - True/False/Not Given: qtype 'mcq', qformat 'tfng', options exactly "
    "['TRUE', 'FALSE', 'NOT GIVEN'], correct_index from the key. Yes/No/Not Given: "
    "the same with qformat 'ynng' and options ['YES', 'NO', 'NOT GIVEN'].\n"
    "   - Multiple choice with ONE answer: qtype 'mcq', no qformat; every option in "
    "`options`, 0-based `correct_index`.\n"
    "   - 'Choose TWO/THREE letters' (several answers from one list): qtype 'mcq', "
    "qformat 'multi_select', all options in `options`, ALL correct 0-based indices in "
    "`correct_indices`, and `select_count` = how many must be chosen. Output the task "
    "ONCE as a single question — do not repeat it per answer.\n"
    "   - Gap-fill / sentence / summary / note / table completion where students write "
    "words FROM THE PASSAGE: qtype 'short', qformat 'gap_fill'. Put the sentence "
    "containing the blank in `prompt`, writing the blank as underscores like '________'. "
    "All acceptable answers (synonyms, British/American spellings) go in `accept_answers`.\n"
    "   - Summary/note completion where students choose from a GIVEN list of words or "
    "phrases (e.g. 'Complete the summary using the list of words, A-D, below'): this is a "
    "matching task, NOT a gap-fill — qtype 'mcq', qformat 'matching', `options` = the "
    "full word list in its original order, correct_index from the key.\n"
    "   - Other short answers (e.g. 'answer in NO MORE THAN TWO WORDS' questions): "
    "qtype 'short', no qformat, answers in `accept_answers`.\n"
    "   - 'explain' for Writing and Speaking tasks and any essay/extended answer. "
    "EVERY question in a writing or speaking section must be 'explain'.\n"
    "4. Set `sub_skill` for every reading/listening question, choosing the single best "
    f"fit from this exact list (no other values): {', '.join(SUB_SKILLS)}.\n"
    "5. COMPLETENESS IS CRITICAL. Before extracting, find the highest question number "
    "printed in the material and put it in source_question_count (official numbering: a "
    "'Choose TWO letters' task covering questions 24-25 counts as TWO; a full IELTS "
    "paper is 40). Then output EVERY question — never skip, merge or summarise "
    "numbered items, even when they look repetitive, span page breaks, sit inside "
    "tables/note templates, or continue after an instruction box. If a question is "
    "partly illegible, still output it with your best reading of the prompt.\n"
    "5b. Keep questions in their original order. Number nothing in the prompt text "
    "itself — the position is enough.\n"
    "5c. Set `task_instructions` on EVERY question: the exact instruction block printed "
    "above its task, word limits included (e.g. 'Complete the notes below. Write ONE WORD "
    "ONLY from the passage for each answer.' or 'Do the following statements agree with "
    "the information given in the text? Choose TRUE, FALSE or NOT GIVEN.'). Every "
    "question belonging to the same task must carry the IDENTICAL string — the test "
    "screen groups consecutive questions that share it into one task box, exactly like "
    "the official computer-based test.\n"
    "6. Answers: if a separate ANSWER SHEET document is supplied, it is the authoritative "
    "source — match its answers to questions by question number and make every "
    "correct_index / accept_answers agree with it exactly, even if you would have answered "
    "differently from the test alone. If there is no answer sheet but a key is printed "
    "inside the test material, use that. If no key exists anywhere, leave "
    "correct_index/accept_answers empty rather than guessing.\n"
    "7. Preserve the original wording of prompts and options; fix only obvious OCR errors.\n"
    "8. MATCHING tasks (matching headings, matching information/features/endings, and any "
    "task where students pick from ONE shared list such as a box of headings i-x or a list "
    "of names): output EACH item as an 'mcq' question with qformat 'matching' whose "
    "`options` are the FULL shared list, in the original order (the site renders them as "
    "a dropdown). Set sub_skill to 'matching_headings'. Put the shared list (e.g. 'List "
    "of Headings') in passage_md as well so nothing is lost.\n"
    "9. Capture EVERYTHING the student needs to answer. Include in passage_md (using simple "
    "markdown) every piece of student-facing material that is not itself a numbered "
    "question: task instructions and word limits, lists of headings, word banks / boxes of "
    "options, example answers, and the text of any notes/table/flow-chart/summary-completion "
    "templates and diagram labels. Reproduce tables as markdown tables and keep gaps as "
    "blanks like '________ (3)'. Never omit a heading list, option box, or instruction line."
)

# Optional enrichment: per-question explanations roughly double the output the
# model must write, which is what makes imports slow. Off by default — the same
# content is generated on demand (and cached) by POST /api/questions/{id}/explain
# the first time a student opens an explanation.
ENRICH_RULE = (
    "\n10. For every reading and listening question, also fill `explanation` (2-3 plain "
    "sentences on why the correct answer is correct, noting why a common wrong choice is a "
    "trap), `support_sentences` (the exact sentence(s) copied verbatim from passage_md "
    "that justify the answer) and `paraphrases` (1-3 pairs showing how the question "
    "rephrases the passage: question_phrase from the question/correct option, "
    "passage_phrase copied VERBATIM from passage_md). Leave all three empty for "
    "writing/speaking questions."
)

_ENRICH_FIELDS = ("explanation", "support_sentences", "paraphrases")


def _lean_tool(tool: dict) -> dict:
    """A copy of TEST_TOOL without the enrichment fields, so a fast import's
    schema doesn't invite the model to write them anyway."""
    import copy
    lean = copy.deepcopy(tool)
    q_props = lean["input_schema"]["properties"]["sections"]["items"][
        "properties"]["questions"]["items"]["properties"]
    for f in _ENRICH_FIELDS:
        q_props.pop(f, None)
    return lean


TEST_TOOL_LEAN = _lean_tool(TEST_TOOL)


def _to_gemini_schema(node):
    """Convert the JSON-schema-ish TEST_TOOL schema to the subset Gemini's
    response_schema accepts (notably, type names must be uppercase)."""
    if not isinstance(node, dict):
        return node
    out = {}
    for key, value in node.items():
        if key == "type" and isinstance(value, str):
            out[key] = value.upper()
        elif key == "properties" and isinstance(value, dict):
            out[key] = {k: _to_gemini_schema(v) for k, v in value.items()}
        elif key == "items":
            out[key] = _to_gemini_schema(value)
        else:
            out[key] = value
    return out


_GEMINI_SCHEMA = _to_gemini_schema(TEST_TOOL["input_schema"])
_GEMINI_SCHEMA_LEAN = _to_gemini_schema(TEST_TOOL_LEAN["input_schema"])




def _extract_text(filename: str, data: bytes, content_type: str = "") -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader  # lazy
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if name.endswith(".docx"):
        from docx import Document  # lazy (python-docx)
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        # Word tests often keep questions/answer keys in tables — don't drop them.
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()
    if name.endswith(".doc"):
        raise HTTPException(
            400,
            "Legacy .doc files aren't supported. Re-save the file as .docx or PDF and try again.",
        )
    if name.endswith((".txt", ".md")) or (content_type or "").startswith("text/"):
        return data.decode("utf-8", errors="replace").strip()
    return ""


def _is_image(name: str, ctype: str) -> bool:
    return ctype.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif"))


def _finalize(result: dict) -> dict:
    """Normalise the model output so the builder always receives a consistent
    shape, regardless of provider or model drift."""
    result.setdefault("difficulty", "medium")
    if result.get("difficulty") not in ("low", "medium", "high"):
        result["difficulty"] = "medium"
    result.setdefault("time_limit_min", 60)

    for sec in result.get("sections") or []:
        skill = (sec.get("skill") or "reading").lower()
        sec["skill"] = skill if skill in ("reading", "listening", "writing", "speaking") else "reading"
        productive = sec["skill"] in ("writing", "speaking")
        for q in sec.get("questions") or []:
            qtype = (q.get("qtype") or "short").lower()
            # Writing/Speaking are always extended, manually-marked answers.
            if productive:
                qtype = "explain"
            elif qtype not in ("mcq", "short", "explain"):
                qtype = "short"
            q["qtype"] = qtype

            # Display format: keep it coherent with the question type.
            qformat = (q.get("qformat") or "").lower() or None
            if qformat not in ("tfng", "ynng", "matching", "multi_select", "gap_fill"):
                qformat = None
            if qtype == "short" and qformat not in (None, "gap_fill"):
                qformat = None
            if qtype == "mcq" and qformat == "gap_fill":
                qformat = None

            if qtype == "mcq":
                opts = [o for o in (q.get("options") or []) if str(o).strip()]
                # TFNG/YNNG always use the fixed three options, whatever the model sent.
                if qformat == "tfng":
                    opts = ["TRUE", "FALSE", "NOT GIVEN"]
                elif qformat == "ynng":
                    opts = ["YES", "NO", "NOT GIVEN"]
                q["options"] = opts
                if qformat == "multi_select":
                    idxs = sorted({
                        int(i) for i in (q.get("correct_indices") or [])
                        if isinstance(i, (int, float)) and 0 <= int(i) < len(opts)
                    })
                    if len(idxs) < 2:
                        # Not really a multi-select — fall back to plain MCQ.
                        qformat = None
                    else:
                        q["correct_indices"] = idxs
                        sc = q.get("select_count")
                        q["select_count"] = sc if isinstance(sc, int) and sc >= 2 else len(idxs)
                        q["correct_index"] = idxs[0]  # harmless fallback for old clients
                if qformat != "multi_select":
                    q.pop("correct_indices", None)
                    q.pop("select_count", None)
                    ci = q.get("correct_index")
                    ok = isinstance(ci, int) and 0 <= ci < len(opts)
                    # Don't hide a missing answer behind a silent default — the
                    # builder shows these questions as "needs answer".
                    q["answer_missing"] = not ok
                    q["correct_index"] = ci if ok else 0
                else:
                    q["answer_missing"] = False
            else:
                q.pop("options", None)
                q.pop("correct_index", None)
                q.pop("correct_indices", None)
                q.pop("select_count", None)
                if qtype == "short":
                    accepts = [str(a).strip() for a in (q.get("accept_answers") or []) if str(a).strip()]
                    q["accept_answers"] = accepts
                    q["answer_missing"] = not accepts
            q["qformat"] = qformat

            # Task grouping: keep the instruction block as a plain stripped
            # string (None when absent) — the test screen groups consecutive
            # questions sharing the identical string into one task box.
            ti = q.get("task_instructions")
            q["task_instructions"] = ti.strip() if isinstance(ti, str) and ti.strip() else None

            # Keep sub_skill within the closed analytics vocab.
            if qtype == "explain":
                q.pop("sub_skill", None)
                q.pop("explanation", None)
                q.pop("support_sentences", None)
                q.pop("paraphrases", None)
            else:
                if q.get("sub_skill") not in SUB_SKILLS:
                    q["sub_skill"] = "multiple_choice" if qtype == "mcq" else "short_answer"
                # Normalise the learning-feature fields.
                exp = q.get("explanation")
                q["explanation"] = exp.strip() if isinstance(exp, str) and exp.strip() else None
                supp = q.get("support_sentences")
                q["support_sentences"] = (
                    [str(s).strip() for s in supp if str(s).strip()]
                    if isinstance(supp, list) else None
                )
                paras = q.get("paraphrases")
                q["paraphrases"] = (
                    [
                        {
                            "question_phrase": str(p.get("question_phrase") or "").strip(),
                            "passage_phrase": str(p.get("passage_phrase") or "").strip(),
                        }
                        for p in paras
                        if isinstance(p, dict)
                        and str(p.get("question_phrase") or "").strip()
                        and str(p.get("passage_phrase") or "").strip()
                    ][:3]
                    if isinstance(paras, list) else None
                ) or None
    return result


def _call_with_retries(fn):
    """Call an LLM request `fn`, retrying transient errors (429/503/529/overloaded)
    with backoff. Raises a friendly HTTPException on give-up."""
    last_err = None
    for attempt in range(4):
        try:
            return fn()
        except HTTPException:
            raise
        except Exception as e:  # noqa: BLE001
            last_err = e
            if any(k in str(e).lower() for k in _TRANSIENT_KEYS) and attempt < 3:
                time.sleep(1.5 * (attempt + 1))
                continue
            break
    m = str(last_err).lower()
    if any(k in m for k in _TRANSIENT_KEYS):
        raise HTTPException(
            503,
            "The AI service is busy right now (high demand). Please wait a few "
            "seconds and try Convert again.",
        )
    raise HTTPException(502, f"AI service error: {last_err}")


def _loads_json(raw: str) -> dict:
    """Parse model JSON, tolerating ```json fences or stray text around the object."""
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw[:4].lower() == "json":
            raw = raw[4:]
        raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        i, j = raw.find("{"), raw.rfind("}")
        if i != -1 and j > i:
            return json.loads(raw[i:j + 1])
        raise


def _gemini_truncated(resp) -> bool:
    """True if Gemini stopped because it hit the output-token ceiling."""
    try:
        fr = resp.candidates[0].finish_reason
    except (AttributeError, IndexError, TypeError):
        return False
    return str(getattr(fr, "name", fr) or "").upper() == "MAX_TOKENS"


def _validate_usable(result: dict) -> dict:
    """Reject an empty/unusable parse with a clear message instead of handing the
    builder a blank test. Fills a default name if the model omitted one, and adds
    an `import_summary` so the teacher can verify completeness and answers."""
    if not (str(result.get("name") or "").strip()):
        result["name"] = "Imported test"
    sections = result.get("sections") or []
    if not sections or not any(sec.get("questions") for sec in sections):
        raise HTTPException(
            422,
            "The AI couldn't find any test questions in that file. Check that you "
            "uploaded the test paper itself (not just an answer sheet) and that the "
            "pages are legible.",
        )

    # Count official question numbers, not rows: a multi-select "Choose N"
    # question occupies N numbers (24-26) on the paper, so the total lines up
    # with source_question_count and a full test reads 40.
    total = 0
    missing = []  # official question-number labels with no answer from the paper/key
    for sec in sections:
        for q in sec.get("questions") or []:
            width = (
                q.get("select_count") or len(q.get("correct_indices") or []) or 2
            ) if q.get("qformat") == "multi_select" else 1
            start = total + 1
            total += width
            if q.get("answer_missing"):
                missing.append(f"{start}" if width == 1 else f"{start}–{total}")
    src = result.pop("source_question_count", None)
    result["import_summary"] = {
        "total_questions": total,
        "source_question_count": src if isinstance(src, int) and src > 0 else None,
        "missing_answers": missing,
    }
    return result


# --- Providers -------------------------------------------------------------

def _gemini_file_parts(types, file: UploadFile, data: bytes, label: str) -> list:
    """Build Gemini content parts for one uploaded file: raw bytes for images and
    PDFs (multimodal — no OCR), locally-extracted text for everything else."""
    name = (file.filename or "").lower()
    ctype = file.content_type or ""
    if _is_image(name, ctype):
        media_type = ctype if ctype.startswith("image/") else "image/png"
        return [types.Part.from_bytes(data=data, mime_type=media_type)]
    if ctype == "application/pdf" or name.endswith(".pdf"):
        return [types.Part.from_bytes(data=data, mime_type="application/pdf")]
    try:
        text = _extract_text(file.filename or "", data, ctype)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"Could not read '{file.filename}': {e}")
    if not text:
        raise HTTPException(
            400,
            f"No readable text found in '{file.filename}'. For scanned/handwritten "
            "material, upload a PDF or image instead.",
        )
    return [types.Part.from_text(text=f"{label}:\n\n{text}")]


ANSWER_SHEET_NOTE = (
    "The next document is the ANSWER SHEET (answer key) for the test above. It is "
    "the authoritative source of answers: match each answer to its question number "
    "and fill correct_index / accept_answers from it exactly."
)


def _run_gemini(file: UploadFile, data: bytes, answers: tuple[UploadFile, bytes] | None = None,
                enrich: bool = False) -> dict:
    if not os.getenv("GEMINI_API_KEY"):
        raise HTTPException(
            503,
            "AI import is not configured. Add GEMINI_API_KEY to backend/.env "
            "(get a free key at aistudio.google.com/apikey), then restart the backend.",
        )
    try:
        from google import genai  # lazy
        from google.genai import types
    except ImportError:
        raise HTTPException(
            503,
            "AI import dependencies are missing. Run: pip install -r backend/requirements.txt",
        )

    parts = _gemini_file_parts(types, file, data, "Test material")
    if answers is not None:
        parts.append(types.Part.from_text(text=ANSWER_SHEET_NOTE))
        parts.extend(_gemini_file_parts(types, answers[0], answers[1], "Answer sheet"))
    parts.append(types.Part.from_text(text="Convert this IELTS test into the structured JSON format."))

    try:
        client = genai.Client(
            api_key=os.getenv("GEMINI_API_KEY"),
            http_options=types.HttpOptions(timeout=int(LLM_TIMEOUT_SECONDS * 1000)),
        )
    except Exception:  # noqa: BLE001 — older SDKs may not accept http_options
        client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
    model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
    cfg_kwargs = dict(
        system_instruction=SYSTEM + (ENRICH_RULE if enrich else ""),
        response_mime_type="application/json",
        response_schema=_GEMINI_SCHEMA if enrich else _GEMINI_SCHEMA_LEAN,
        temperature=0,  # deterministic: same file -> same structured test
        max_output_tokens=GEMINI_MAX_OUTPUT_TOKENS,  # full tests are long — avoid truncation
    )
    # Import is pure extraction — the model copies questions and matches answers
    # from the key (it never solves the test; rule 6 forbids guessing). Gemini
    # 2.5's default "thinking" phase adds nothing here but more than doubles the
    # import time, so it's off unless GEMINI_THINKING_BUDGET says otherwise.
    try:
        cfg = types.GenerateContentConfig(
            **cfg_kwargs,
            thinking_config=types.ThinkingConfig(
                thinking_budget=int(os.getenv("GEMINI_THINKING_BUDGET", "0")),
            ),
        )
    except (TypeError, ValueError):  # older SDKs / models without thinking
        cfg = types.GenerateContentConfig(**cfg_kwargs)

    # Gemini's free tier returns transient 503/429 spikes — retry with backoff.
    resp = _call_with_retries(
        lambda: client.models.generate_content(model=model, contents=parts, config=cfg)
    )

    if _gemini_truncated(resp):
        raise HTTPException(
            502,
            "That test is too large to import in one go. Split it into separate "
            "reading/listening files and import each, then combine in the builder.",
        )

    raw = (resp.text or "").strip()
    if not raw:
        raise HTTPException(502, "The AI did not return a usable test. Try a clearer file.")
    try:
        result = _loads_json(raw)
    except json.JSONDecodeError:
        raise HTTPException(502, "The AI returned malformed data. Try again or use a clearer file.")
    return _validate_usable(_finalize(result))


def _claude_file_blocks(file: UploadFile, data: bytes, label: str) -> list:
    """Build Claude content blocks for one uploaded file: base64 image/document
    blocks for images and PDFs (so scans work), extracted text for the rest."""
    name = (file.filename or "").lower()
    ctype = file.content_type or ""
    if _is_image(name, ctype):
        media_type = ctype if ctype.startswith("image/") else "image/png"
        return [{
            "type": "image",
            "source": {"type": "base64", "media_type": media_type,
                       "data": base64.standard_b64encode(data).decode()},
        }]
    if ctype == "application/pdf" or name.endswith(".pdf"):
        return [{
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf",
                       "data": base64.standard_b64encode(data).decode()},
        }]
    try:
        text = _extract_text(file.filename or "", data, ctype)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"Could not read '{file.filename}': {e}")
    if not text:
        raise HTTPException(
            400,
            f"No readable text found in '{file.filename}'. For scanned/handwritten "
            "material, upload a PDF or image instead.",
        )
    return [{"type": "text", "text": f"{label}:\n\n{text}"}]


def _run_claude(file: UploadFile, data: bytes, answers: tuple[UploadFile, bytes] | None = None,
                enrich: bool = False) -> dict:
    if not os.getenv("ANTHROPIC_API_KEY"):
        raise HTTPException(
            503,
            "AI import is not configured. Add ANTHROPIC_API_KEY to backend/.env "
            "(get a key at console.anthropic.com), then restart the backend.",
        )
    try:
        import anthropic  # lazy
    except ImportError:
        raise HTTPException(
            503,
            "AI import dependencies are missing. Run: pip install -r backend/requirements.txt",
        )

    user_content = _claude_file_blocks(file, data, "Test material")
    if answers is not None:
        user_content.append({"type": "text", "text": ANSWER_SHEET_NOTE})
        user_content.extend(_claude_file_blocks(answers[0], answers[1], "Answer sheet"))
    user_content.append({"type": "text", "text": "Convert this IELTS test into the structured format."})

    client = anthropic.Anthropic(timeout=LLM_TIMEOUT_SECONDS)
    resp = _call_with_retries(lambda: client.messages.create(
        model=MODEL,
        max_tokens=CLAUDE_MAX_OUTPUT_TOKENS,
        system=[{"type": "text", "text": SYSTEM + (ENRICH_RULE if enrich else ""),
                 "cache_control": {"type": "ephemeral"}}],
        tools=[TEST_TOOL if enrich else TEST_TOOL_LEAN],
        tool_choice={"type": "tool", "name": "build_ielts_test"},
        messages=[{"role": "user", "content": user_content}],
    ))

    if getattr(resp, "stop_reason", None) == "max_tokens":
        raise HTTPException(
            502,
            "That test is too large to import in one go. Split it into separate "
            "files and import each, then combine in the builder.",
        )

    for block in resp.content:
        if getattr(block, "type", None) == "tool_use" and block.name == "build_ielts_test":
            return _validate_usable(_finalize(block.input))

    raise HTTPException(502, "The AI did not return a usable test. Try a clearer file.")


def _local_text(file: UploadFile, data: bytes) -> str:
    name = (file.filename or "").lower()
    ctype = file.content_type or ""
    if _is_image(name, ctype):
        raise HTTPException(
            400,
            "The local importer cannot read images. Set LLM_PROVIDER=gemini for "
            "image/scanned tests, or upload a PDF/Word file.",
        )
    try:
        text = _extract_text(file.filename or "", data, ctype)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"Could not read '{file.filename}': {e}")
    if not text:
        raise HTTPException(
            400,
            "No readable text found. Set LLM_PROVIDER=gemini to read scanned/image tests.",
        )
    return text


def _run_local(file: UploadFile, data: bytes, answers: tuple[UploadFile, bytes] | None = None,
               enrich: bool = False) -> dict:
    """No-API fallback: extract text and drop it into one reading section so the
    teacher can build questions by hand. An answer sheet, if given, is appended
    so the teacher can copy answers in while building."""
    text = _local_text(file, data)
    if answers is not None:
        text += "\n\n---\n\n## Answer sheet\n\n" + _local_text(answers[0], answers[1])
    return _finalize({
        "name": file.filename or "Imported test",
        "sections": [{
            "skill": "reading",
            "title": "Imported material",
            "passage_md": text,
            "questions": [],
        }],
    })


_PROVIDERS = {"gemini": _run_gemini, "claude": _run_claude, "local": _run_local}


async def _read_upload(file: UploadFile) -> bytes:
    # Read one byte past the limit so we can detect (and reject) oversized files
    # without loading the whole upload into memory.
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            413,
            f"'{file.filename}' is too large. The maximum upload size is "
            f"{MAX_UPLOAD_BYTES // (1024 * 1024)} MB.",
        )
    if not data:
        raise HTTPException(400, f"The uploaded file '{file.filename}' is empty.")
    return data


@router.post("/ai")
async def ai_import(
    file: UploadFile = File(...),
    answer_sheet: UploadFile | None = File(None),
    # Also generate per-question explanations during the import. Roughly doubles
    # the AI output (= import time); off by default because explanations are
    # generated on demand (and cached) the first time a student opens one.
    enrich: bool = Form(False),
    db: Session = Depends(get_db),
    user: models.User = Depends(require_role("teacher", "admin")),
    _rl: None = Depends(_ai_limiter),
):
    provider = os.getenv("LLM_PROVIDER", "gemini").lower()
    run = _PROVIDERS.get(provider)
    if run is None:
        raise HTTPException(
            500,
            f"Unknown LLM_PROVIDER '{provider}'. Use 'gemini', 'claude', or 'local'.",
        )

    data = await _read_upload(file)
    answers = None
    if answer_sheet is not None and (answer_sheet.filename or "").strip():
        answers = (answer_sheet, await _read_upload(answer_sheet))
    return run(file, data, answers, enrich=enrich)
